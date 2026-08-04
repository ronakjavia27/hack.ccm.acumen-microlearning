"""
flashcards.py - Shared helpers for the flashcard generator pipeline.

Handles:
- Source file parsing (md/txt/pdf/docx/html) -> (title, text)
- LLM conversion prompt + markdown/JSON result parsing
- Sub-topic tagging of decks against the subtopics.txt vocabulary
- Ledger (flashcards_ledger.json) + sha256 helpers for skip-if-unchanged
"""

import hashlib
import json
import os
import re
from datetime import datetime
from uuid import uuid4

from .config import (
    FLASHCARDS_DIR,
    FLASHCARDS_INDEX_FILE,
    FLASHCARDS_LEDGER_FILE,
    THEORY_SPEC_TO_CANONICAL,
    MAX_CARDS_PER_DECK,
)
from .subtopics_config import get_subtopics_for_system
from .tracking import save_json_atomic
SUPPORTED_EXTS = (".md", ".txt", ".pdf", ".docx", ".html", ".htm")

# =====================================================================
# SOURCE FILE PARSING
# =====================================================================

def _read_utf8(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _title_from_markdown(text, stem):
    for line in text.split("\n")[:30]:
        s = line.strip()
        if s.startswith("# "):
            return s[2:].strip()
    m = re.match(r"^\*\*(\d+\.?\s*)?([^*]+?)\*\*\s*:?\s*$", text.split("\n")[0].strip() if text.split("\n") else "")
    if m and m.group(2):
        return m.group(2).strip()
    return re.sub(r"^\d+\s*", "", stem).strip()


def extract_md_txt(path):
    text = _read_utf8(path)
    stem = os.path.splitext(os.path.basename(path))[0]
    return _title_from_markdown(text, stem), text.strip()


def extract_pdf(path, ocr_fallback=True):
    from pypdf import PdfReader
    reader = PdfReader(path)
    chunks = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(chunks).strip()
    if ocr_fallback and len(text) < 150:
        try:
            from .ocr import fallback_page_ocr
            text = fallback_page_ocr(path)
        except Exception:
            pass
    stem = os.path.splitext(os.path.basename(path))[0]
    return re.sub(r"^\d+\s*", "", stem).strip(), text


def extract_docx(path):
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is required for .docx files. Install it with: pip install python-docx"
        )
    doc = Document(path)
    parts = []
    for block in doc.paragraphs:
        t = block.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            parts.append(" | ".join(cells))
    stem = os.path.splitext(os.path.basename(path))[0]
    return re.sub(r"^\d+\s*", "", stem).strip(), "\n".join(parts)


def extract_html(path):
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(_read_utf8(path), "html.parser")
    title = ""
    if soup.title and soup.title.get_text(strip=True):
        title = soup.title.get_text(strip=True)
    elif soup.h1 and soup.h1.get_text(strip=True):
        title = soup.h1.get_text(strip=True)
    text = soup.get_text("\n", strip=True)
    stem = os.path.splitext(os.path.basename(path))[0]
    return title or re.sub(r"^\d+\s*", "", stem).strip(), text


def parse_source_file(path):
    """Parse a source file -> (title, text). Raises ValueError for unsupported types."""
    ext = os.path.splitext(str(path))[1].lower()
    if ext in (".md", ".txt"):
        return extract_md_txt(path)
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext in (".html", ".htm"):
        return extract_html(path)
    raise ValueError(f"Unsupported file type: {ext}")


# =====================================================================
# LLM CONVERSION (source -> markdown deck)
# =====================================================================

CONVERT_SYSTEM_PROMPT = """You convert medical study material into a markdown flashcard deck.
Rules:
- Each flashcard starts with a line "## " followed by a short card title (one card per section).
- Card body: concise, high-yield study notes. Use tables, bullet lists, and bold key terms where useful. Refine, summarise, and reformat the material — do not copy it verbatim.
- Keep exact numbers, doses, thresholds, and units — never invent or change clinical values. If the source omits a value, omit it too.
- Optional: put a "---" divider inside a card to separate a front (question/prompt) from the back (answer) for flip-mode study. Most cards can omit it.
- Generate 5-12 cards per source chunk (maximum """ + str(MAX_CARDS_PER_DECK) + """).
- Output ONLY the markdown. No preamble, no code fences, no commentary."""


def _extract_markdown(result):
    """Normalize LLM output (str / JSON dict) into a markdown string."""
    if isinstance(result, str):
        return result.strip()
    if isinstance(result, dict):
        for key in ("markdown", "content", "output", "deck", "text"):
            v = result.get(key)
            if isinstance(v, str) and v.strip():
                return v.strip()
        cards = result.get("cards")
        if isinstance(cards, list) and cards:
            parts = []
            for c in cards:
                if not isinstance(c, dict):
                    continue
                parts.append("## " + str(c.get("subtopic", "Card")).strip() + "\n\n" + str(c.get("content", "")).strip())
            return "\n\n".join(parts)
        for v in result.values():
            if isinstance(v, str) and v.strip():
                return v.strip()
    return ""


def _cards_from_dict(obj, title):
    cards = obj.get("cards") if isinstance(obj, dict) else None
    if isinstance(cards, list) and cards:
        out = []
        for c in cards:
            if not isinstance(c, dict):
                continue
            body = str(c.get("content", "")).strip()
            if not body:
                continue
            out.append({"subtopic": str(c.get("subtopic", "Card")).strip(), "content": body})
        return out
    return []


def _cards_from_markdown(markdown_text, title):
    lines = markdown_text.split("\n")
    card_starts = []
    for i, line in enumerate(lines):
        if line.strip().startswith("## "):
            card_starts.append((i, line.strip()[3:].strip()))
    if not card_starts:
        body = markdown_text.strip()
        if body:
            return [{"subtopic": title, "content": body}]
        return []
    cards = []
    for idx, (start_line, heading) in enumerate(card_starts):
        end_line = card_starts[idx + 1][0] if idx + 1 < len(card_starts) else len(lines)
        body = "\n".join(lines[start_line + 1:end_line]).strip()
        if not body:
            continue
        cards.append({"subtopic": heading, "content": body})
    return cards


def parse_llm_markdown(result, title):
    """Parse LLM output into a list of card dicts {subtopic, content}.
    Accepts markdown text, a JSON string ({"cards": [...]} or {"markdown": "..."}),
    or a parsed dict."""
    if isinstance(result, dict):
        md = _extract_markdown(result)
        if md:
            return _cards_from_markdown(md, title)
        return _cards_from_dict(result, title)
    if not isinstance(result, str):
        return []
    stripped = result.strip()
    if stripped.startswith(("{", "[")):
        try:
            obj = json.loads(stripped)
            if isinstance(obj, dict):
                md = _extract_markdown(obj)
                if md and "cards" not in obj:
                    return _cards_from_markdown(md, title)
                return _cards_from_dict(obj, title)
        except json.JSONDecodeError:
            pass
        m = re.search(r'(\{.*"cards".*\})', stripped, re.DOTALL)
        if m:
            try:
                return _cards_from_dict(json.loads(m.group(1)), title)
            except json.JSONDecodeError:
                pass
    return _cards_from_markdown(stripped, title)


def llm_convert_to_markdown(text, title, llm="openrouter", verbose=False, api_key=None, model=None, base_url=None):
    """Send source text to the LLM, return markdown deck text (or None on failure).

    When api_key or model is given, a direct OpenAI-compatible call is used
    (base_url defaults to OpenRouter). Otherwise the default openrouter path
    uses the dedicated flashcard model (FLASHCARD_LLM_*)."""
    from .llm import execute_openai_compat, execute_with_fallback, execute_with_gemini

    user_prompt = (
        f"Source title: {title}\n\nSource content:\n{text}\n\n"
        "Generate the flashcard deck in markdown following the rules."
    )
    try:
        if api_key or model:
            result = execute_openai_compat(
                CONVERT_SYSTEM_PROMPT, user_prompt,
                api_key=api_key, model=model, base_url=base_url,
                json_mode=False,
            )
        elif llm == "gemini":
            result = execute_with_gemini(CONVERT_SYSTEM_PROMPT, user_prompt, "flashcards")
        elif llm == "together":
            result = execute_with_fallback(CONVERT_SYSTEM_PROMPT, user_prompt, "flashcards")
        else:
            result = flashcard_llm(CONVERT_SYSTEM_PROMPT, user_prompt)
        return _extract_markdown(result) or None
    except Exception as e:
        if verbose:
            print(f"    [X] LLM convert failed: {e}")
        return None


# =====================================================================
# SUBTOPIC TAGGING (subtopics.txt vocabulary)
# =====================================================================

TAG_SYSTEM_PROMPT = """You assign subtopics to medical flashcard cards. Use ONLY the exact subtopics from the provided vocabulary — do not invent or rephrase them.
For each card index, pick 1-3 of the valid subtopics that best match the card content.
Return ONLY valid JSON: {{"cards": [{{"index": 0, "tags": ["Exact Subtopic"]}}, ...]}}. Include every card index exactly once. No preamble, no commentary."""


def canonical_systems_for(specialty):
    """THEORY folder abbreviation -> canonical subtopic-vocab systems."""
    return THEORY_SPEC_TO_CANONICAL.get(str(specialty).strip()) or []


def _acronym(phrase):
    words = re.findall(r"[A-Za-z0-9]+", str(phrase))
    if len(words) == 1:
        return words[0].lower()
    return "".join(w[0] for w in words).lower()


def normalize_subtopic(system, candidate):
    """Match an LLM-supplied tag against the vocab for a canonical system.
    Returns the exact vocab entry, or None if no match."""
    cand = str(candidate or "").strip().lstrip("#").strip()
    if not cand:
        return None
    subs = get_subtopics_for_system(system) or []
    cl = cand.lower()
    ca = _acronym(cand)
    for s in subs:
        if s.lower() == cl:
            return s
    for s in subs:
        if cl in s.lower() or s.lower() in cl:
            return s
    for s in subs:
        if ca and (_acronym(s) == ca or re.sub(r"[\s\-&]+", "", s.lower()) == re.sub(r"[\s\-&]+", "", cl)):
            return s
    return None


def tag_cards_with_llm(cards, systems, llm="openrouter", verbose=False, api_key=None, model=None, base_url=None):
    """Tag every card dict in `cards` (list) with subtopics from the vocabulary.
    Mutates cards[i]['tags'] in place. systems = canonical system names for the vocab.
    Returns the number of tagged cards.

    When api_key or model is given, a direct OpenAI-compatible call is used
    (base_url defaults to OpenRouter). Otherwise the dedicated flashcard model
    (FLASHCARD_LLM_*) is used, unless llm == 'gemini'/'together' (provider chain)."""
    from .llm import execute_openai_compat, execute_with_fallback, execute_with_gemini
    if not systems:
        if verbose:
            print("    [~] No canonical systems — skipping tags")
        return 0
    vocab = []
    for s in systems:
        for sub in get_subtopics_for_system(s) or []:
            if sub not in vocab:
                vocab.append(sub)
    if not vocab or not cards:
        return 0

    vocab_str = "\n".join(f"{i+1}. {v}" for i, v in enumerate(vocab))
    system_prompt = TAG_SYSTEM_PROMPT.format(vocab_str=vocab_str, systems=", ".join(systems))
    lines = []
    for i, c in enumerate(cards):
        lines.append(f"[{i}] {c.get('subtopic', '')}\n{(c.get('content') or c.get('back') or '')[:800]}")
    user_prompt = "Valid subtopics:\n" + vocab_str + "\n\nCards:\n\n" + "\n\n".join(lines)

    try:
        if api_key or model:
            result = execute_openai_compat(
                system_prompt, user_prompt,
                api_key=api_key, model=model, base_url=base_url,
                json_mode=True,
            )
        elif llm == "gemini":
            result = execute_with_gemini(system_prompt, user_prompt, "flashcards")
        elif llm == "together":
            result = execute_with_fallback(system_prompt, user_prompt, "flashcards")
        else:
            result = flashcard_llm(system_prompt, user_prompt, json_mode=True)
    except Exception as e:
        if verbose:
            print(f"    [X] LLM tagging failed: {e}")
        return 0

    mapping = {}
    if isinstance(result, dict):
        raw_cards = result.get("cards")
    elif isinstance(result, str):
        raw_cards = None
        stripped = result.strip()
        if stripped.startswith(("{", "[")):
            try:
                obj = json.loads(stripped)
                if isinstance(obj, dict):
                    raw_cards = obj.get("cards")
            except json.JSONDecodeError:
                m = re.search(r'(\{.*"cards".*\})', stripped, re.DOTALL)
                if m:
                    try:
                        raw_cards = json.loads(m.group(1)).get("cards")
                    except json.JSONDecodeError:
                        pass
    else:
        raw_cards = None
    for entry in raw_cards or []:
        if not isinstance(entry, dict):
            continue
        try:
            idx = int(entry.get("index"))
        except (TypeError, ValueError):
            continue
        mapping[idx] = entry.get("tags") or []

    tagged = 0
    for i, c in enumerate(cards):
        tags = []
        for cand in mapping.get(i, []):
            norm = None
            for s in systems:
                norm = normalize_subtopic(s, cand)
                if norm:
                    break
            if norm and norm not in tags:
                tags.append(norm)
            if len(tags) >= 3:
                break
        c["tags"] = tags
        if tags:
            tagged += 1
    if verbose:
        print(f"    [i] Tagged {tagged}/{len(cards)} cards")
    return tagged


# =====================================================================
# FRONT QUESTION GENERATION (portal flip mode)
# =====================================================================

QUESTION_SYSTEM_PROMPT = (
    "You write the front question of a medical flashcard. Given a deck "
    "title, a card subtopic, and the card's answer content, produce ONE "
    "short, crisp question (max ~15 words) a clinician would ask to recall "
    "exactly that content. The question must be answerable by the card "
    "content alone. Do not include the answer, numbering, quotes, markdown, "
    "or any explanation. Output only the question text."
)


def generate_card_question(deck_title, subtopic, content, api_key=None, model=None, base_url=None, verbose=False):
    """Generate a short flashcard front question.

    Uses the QUESTION_LLM_* config — a deliberately different model/provider
    than every other pipeline. api_key/model/base_url overrides point at an
    OpenAI-compatible endpoint (default: OpenRouter + QUESTION_LLM_MODEL).
    Returns the question string or None on failure."""
    from .config import (
        QUESTION_LLM_API_KEY,
        QUESTION_LLM_MODEL,
        QUESTION_LLM_BASE_URL,
        TEMPERATURE_QUESTION,
        MAX_TOKENS_QUESTION,
    )
    from .llm import execute_openai_compat

    key = api_key or QUESTION_LLM_API_KEY
    mdl = model or QUESTION_LLM_MODEL
    endpoint = base_url or QUESTION_LLM_BASE_URL
    if not key:
        return None

    snippet = (content or "").strip()
    if len(snippet) > 1500:
        snippet = snippet[:1500] + "..."
    user_prompt = (
        f"Deck title: {deck_title or ''}\n"
        f"Card subtopic: {subtopic or ''}\n"
        f"Card content:\n{snippet}"
    )
    try:
        result = execute_openai_compat(
            QUESTION_SYSTEM_PROMPT,
            user_prompt,
            api_key=key,
            model=mdl,
            base_url=endpoint,
            temperature=TEMPERATURE_QUESTION,
            max_tokens=MAX_TOKENS_QUESTION,
            json_mode=False,
        )
    except Exception as e:
        if verbose:
            print(f"    [X] Question generation failed: {e}")
        return None

    if isinstance(result, dict):
        text = str(result.get("question") or result.get("text") or result)
    else:
        text = str(result)
    text = text.strip().strip('"\'')
    text = re.sub(r"^(Question|Q)\s*[:\-]\s*", "", text, flags=re.I).strip()
    return text or None


# =====================================================================
# LEDGER (skip-if-unchanged)
# =====================================================================

def file_sha256(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def load_ledger():
    if not os.path.exists(FLASHCARDS_LEDGER_FILE):
        return {}
    try:
        with open(FLASHCARDS_LEDGER_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except (json.JSONDecodeError, Exception):
        return {}


def save_ledger(ledger):
    save_json_atomic(FLASHCARDS_LEDGER_FILE, ledger)


# =====================================================================
# DEDICATED FLASHCARD LLM (FLASHCARD_LLM_* config)
# =====================================================================

def flashcard_llm(system_prompt, user_prompt, json_mode=False, verbose=False, max_tokens=None):
    """Run an LLM call through the dedicated flashcard-processing model
    (FLASHCARD_LLM_MODEL/FLASHCARD_LLM_API_KEY/FLASHCARD_LLM_BASE_URL).
    This is the cheap convert/tag/regenerate layer — separate from the
    QUESTION_LLM_* front-question model and from the main pipeline chain.
    Returns parsed JSON when json_mode=True, else the text string, or None."""
    from .config import (
        FLASHCARD_LLM_API_KEY,
        FLASHCARD_LLM_MODEL,
        FLASHCARD_LLM_BASE_URL,
        MAX_TOKENS_FLASHCARDS,
        TEMPERATURE_FLASHCARDS,
    )
    from .llm import execute_openai_compat
    key = FLASHCARD_LLM_API_KEY
    if not key:
        if verbose:
            print("    [X] No FLASHCARD_LLM_API_KEY — cannot call flashcard LLM")
        return None
    try:
        return execute_openai_compat(
            system_prompt,
            user_prompt,
            api_key=key,
            model=FLASHCARD_LLM_MODEL,
            base_url=FLASHCARD_LLM_BASE_URL,
            temperature=TEMPERATURE_FLASHCARDS,
            max_tokens=max_tokens or MAX_TOKENS_FLASHCARDS,
            json_mode=json_mode,
        )
    except Exception as e:
        if verbose:
            print(f"    [X] flashcard_llm call failed: {e}")
        return None


# =====================================================================
# UNIFIED CARD SCHEMA + STORE (output_files/flashcards/{System}/{Subtopic}.json)
# =====================================================================

def slugify(text):
    """Human-readable slug: lowercase, alphanumeric + hyphens."""
    slug = re.sub(r"[^a-z0-9]+", "-", (text or "").lower()).strip("-")
    return slug or "card"


def build_card(front, back, system, subtopic, tags=None, source="llm", source_file=None,
               data=None, slug_hint=None, card_id=None):
    """Create a store-format card dict with a persistent UUID id."""
    now = datetime.now().isoformat(timespec="seconds")
    return {
        "id": card_id or str(uuid4()),
        "slug": slugify(slug_hint or front or back or subtopic or "card"),
        "front": (front or "").strip(),
        "back": (back or "").strip(),
        "system": system,
        "subtopic": subtopic,
        "tags": tags or [],
        "data": data or {},
        "source": source,
        "source_file": source_file,
        "status": "active",
        "edit_history": [],
        "created_at": now,
        "updated_at": now,
    }


def parse_card_content(content):
    """Split raw card body on an optional '---' divider into (front, back).
    No divider -> (None, content)."""
    if not content:
        return None, ""
    m = re.search(r"\n---\s*\n", content)
    if not m:
        return None, content.strip()
    front = content[:m.start()].strip()
    back = content[m.end():].strip()
    return (front or None), back


def canonical_subtopic(system, candidate):
    """Map a candidate subtopic to the controlled vocabulary for the system.
    Falls back to the candidate (title-cased, trimmed)."""
    norm = normalize_subtopic(system, candidate)
    return norm or (candidate or "").strip()


def _subtopic_slug(system, subtopic):
    return slugify(subtopic or "General")[:60]


def subtopic_file_path(system, subtopic):
    return os.path.join(FLASHCARDS_DIR, system, _subtopic_slug(system, subtopic) + ".json")


def load_subtopic_file(system, subtopic):
    """Read one store file. Returns dict or None."""
    path = subtopic_file_path(system, subtopic)
    if not os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, dict):
            return None
        data.setdefault("id", f"{system}/{subtopic}")
        data.setdefault("system", system)
        data.setdefault("subtopic", subtopic)
        data.setdefault("cards", [])
        data["card_count"] = len(data["cards"])
        return data
    except (json.JSONDecodeError, OSError):
        return None


def write_subtopic_file(data):
    """Atomically write one store file (canonical per (system, subtopic))."""
    system = data.get("system") or "General"
    subtopic = data.get("subtopic") or "General"
    data["id"] = f"{system}/{subtopic}"
    data["card_count"] = len(data.get("cards") or [])
    data["updated_at"] = datetime.now().isoformat(timespec="seconds")
    path = subtopic_file_path(system, subtopic)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    save_json_atomic(path, data)
    rebuild_flashcards_index()


def _save_card(card):
    """Persist a single card into its (system, subtopic) file. Rebuilds index.

    If the card already lives in a different (system, subtopic) file — e.g. its
    subtopic was retagged — the old copy is removed first so every card exists
    in exactly one location."""
    system = card.get("system") or "General"
    subtopic = canonical_subtopic(system, card.get("subtopic"))
    card["system"] = system
    card["subtopic"] = subtopic
    loc = load_flashcards_index().get("card_locations", {}).get(card.get("id"))
    if loc and (loc["system"] != system or loc["subtopic"] != subtopic):
        remove_card_from(card.get("id"), loc["system"], loc["subtopic"], rebuild=False)
    card["subtopic"] = subtopic
    data = load_subtopic_file(system, subtopic)
    if data is None:
        data = {"id": f"{system}/{subtopic}", "system": system,
                "subtopic": subtopic, "cards": [], "card_count": 0,
                "created_at": datetime.now().isoformat(timespec="seconds")}
    cards = data["cards"]
    for i, c in enumerate(cards):
        if c.get("id") == card.get("id"):
            cards[i] = card
            break
    else:
        cards.append(card)
    write_subtopic_file(data)


def upsert_card(card):
    """Insert or update a store card (handles moves: old location card removed)."""
    old_sys = card.pop("_old_system", None)
    old_sub = card.pop("_old_subtopic", None)
    if old_sys and old_sub and (old_sys != card.get("system") or old_sub != card.get("subtopic")):
        remove_card_from(card.get("id"), old_sys, old_sub, rebuild=False)
    card["updated_at"] = datetime.now().isoformat(timespec="seconds")
    _save_card(card)


def remove_card_from(card_id, system, subtopic, rebuild=True):
    """Remove a card by id from a specific store file."""
    data = load_subtopic_file(system, subtopic)
    if data is None:
        return False
    before = len(data["cards"])
    data["cards"] = [c for c in data["cards"] if c.get("id") != card_id]
    if len(data["cards"]) == before:
        return False
    if data["cards"]:
        write_subtopic_file(data)
    else:
        path = subtopic_file_path(system, subtopic)
        try:
            os.remove(path)
        except OSError:
            pass
        if rebuild:
            rebuild_flashcards_index()
    return True


def remove_card(card_id):
    """Remove a card by id wherever it lives (via index lookup)."""
    idx = load_flashcards_index()
    loc = idx.get("card_locations", {}).get(card_id)
    if loc:
        return remove_card_from(card_id, loc["system"], loc["subtopic"])
    return False


def move_card(card_id, new_system, new_subtopic):
    """Relocate a card to another (system, subtopic), updating its fields."""
    idx = load_flashcards_index()
    loc = idx.get("card_locations", {}).get(card_id)
    if not loc:
        return False
    data = load_subtopic_file(loc["system"], loc["subtopic"])
    card = next((c for c in (data or {}).get("cards", []) if c.get("id") == card_id), None)
    if card is None:
        return False
    remove_card_from(card_id, loc["system"], loc["subtopic"], rebuild=False)
    card["system"] = new_system
    card["subtopic"] = canonical_subtopic(new_system, new_subtopic)
    card["updated_at"] = datetime.now().isoformat(timespec="seconds")
    card.setdefault("edit_history", []).append(
        f"moved {loc['system']}/{loc['subtopic']} -> {card['system']}/{card['subtopic']} "
        f"at {datetime.now().isoformat(timespec='seconds')}"
    )
    _save_card(card)
    return True


def find_card(card_id):
    """Locate and return a card dict by UUID, or None."""
    idx = load_flashcards_index()
    loc = idx.get("card_locations", {}).get(card_id)
    if not loc:
        return None
    data = load_subtopic_file(loc["system"], loc["subtopic"])
    return next((c for c in (data or {}).get("cards", []) if c.get("id") == card_id), None)


def rebuild_flashcards_index():
    """Scan the store tree and rebuild the derived master index. Canonical
    data lives in the per-subtopic files; this index is always regenerated."""
    index = {
        "version": 2,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "total_cards": 0,
        "systems": {},
        "cards": {},
        "card_locations": {},
    }
    if not os.path.isdir(FLASHCARDS_DIR):
        save_json_atomic(FLASHCARDS_INDEX_FILE, index)
        return index
    for system in sorted(os.listdir(FLASHCARDS_DIR)):
        sys_dir = os.path.join(FLASHCARDS_DIR, system)
        if not os.path.isdir(sys_dir):
            continue
        index["systems"][system] = {}
        for fname in sorted(os.listdir(sys_dir)):
            if not fname.endswith(".json"):
                continue
            path = os.path.join(sys_dir, fname)
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
            except (json.JSONDecodeError, OSError):
                continue
            subtopic = data.get("subtopic") or os.path.splitext(fname)[0]
            cards = data.get("cards") or []
            index["systems"][system][subtopic] = {"file": fname, "count": len(cards)}
            index["total_cards"] += len(cards)
            for c in cards:
                cid = c.get("id")
                if not cid:
                    continue
                index["cards"][cid] = {
                    "slug": c.get("slug"),
                    "front": (c.get("front") or "")[:300],
                    "system": system,
                    "subtopic": subtopic,
                    "tags": c.get("tags") or [],
                    "source": c.get("source"),
                    "status": c.get("status"),
                    "updated_at": c.get("updated_at"),
                }
                index["card_locations"][cid] = {"system": system, "subtopic": subtopic}
        if not index["systems"][system]:
            del index["systems"][system]
    save_json_atomic(FLASHCARDS_INDEX_FILE, index)
    return index


def load_flashcards_index():
    """Read the master index; regenerate if missing/outdated."""
    try:
        with open(FLASHCARDS_INDEX_FILE, "r", encoding="utf-8") as f:
            idx = json.load(f)
        if isinstance(idx, dict) and idx.get("version") == 2 and "card_locations" in idx:
            return idx
    except (json.JSONDecodeError, OSError):
        pass
    return rebuild_flashcards_index()


def store_cards_all():
    """Yield every (system, subtopic, card) in the store."""
    idx = load_flashcards_index()
    for cid, loc in idx.get("card_locations", {}).items():
        data = load_subtopic_file(loc["system"], loc["subtopic"])
        card = next((c for c in (data or {}).get("cards", []) if c.get("id") == cid), None)
        if card is not None:
            yield loc["system"], loc["subtopic"], card


# =====================================================================
# FRONT QUESTION GENERATION (store-aware batch)
# =====================================================================

def ensure_fronts(cards, verbose=False, api_key=None, model=None, base_url=None, persist=False):
    """Batch-generate front questions for every card missing one (explicit
    front/back cards keep their authored front). Uses the QUESTION_LLM_*
    model — the deliberate separate provider. With persist=True, generated
    fronts are written back to the store files immediately. Returns number generated."""
    missing = [c for c in cards if not (c.get("front") or "").strip()]
    if not missing:
        return 0
    generated = 0
    for c in missing:
        q = generate_card_question(
            c.get("subtopic"), c.get("subtopic"), c.get("back"),
            api_key=api_key, model=model, base_url=base_url, verbose=verbose,
        )
        if q:
            c["front"] = q
            c.setdefault("edit_history", []).append(
                f"front auto-generated at {datetime.now().isoformat(timespec='seconds')}"
            )
            if persist:
                _save_card(c)
            generated += 1
        if verbose:
            print(f"    [i] front {'ok' if q else 'failed'} for card {c.get('id', '?')[:8]}")
    return generated


# =====================================================================
# MARKDOWN DECK IMPORT (authored decks in flashcards_input/)
# =====================================================================

def parse_markdown_deck(md_path):
    """Split an authored markdown deck into (title, [{subtopic, content}]).

    Authoring convention:
        # Deck Title (optional — falls back to filename)
        ## Card title
        ... markdown content (tables, bullets, notes) ...
        ## Another card
        ...
        --- (optional divider inside a card -> front/back faces for flip mode)
    """
    with open(md_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    lines = content.split("\n")

    title = None
    card_starts = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not title and stripped.startswith("# ") and not stripped.startswith("## "):
            title = stripped[2:].strip()
        if stripped.startswith("## "):
            card_starts.append((i, stripped[3:].strip()))

    if not title:
        name = os.path.splitext(os.path.basename(md_path))[0]
        title = re.sub(r"^\d+\s*", "", name).strip()

    if not card_starts:
        body = "\n".join(lines)
        body = re.sub(r"^#\s+.*$", "", body, count=1).strip()
        card_starts = [(0, title)] if body.strip() else []

    cards = []
    for idx, (start_line, heading) in enumerate(card_starts):
        end_line = card_starts[idx + 1][0] if idx + 1 < len(card_starts) else len(lines)
        body = "\n".join(lines[start_line + 1:end_line]).strip()
        if not body:
            continue
        cards.append({"subtopic": heading, "content": body})
    return title, cards


def load_import_meta(source_file):
    """{card_index: existing store card} for a source md, so re-imports keep
    console curation (tags/status/history) AND stable UUIDs by card position.
    The authored md is the source of truth, so card order is stable."""
    source_file = (source_file or "").replace("\\", "/")
    meta = {}
    try:
        for _system, _subtopic, c in store_cards_all():
            if (c.get("source_file") or "").replace("\\", "/") != source_file:
                continue
            idx = len(meta)
            meta[idx] = c
    except Exception:
        pass
    return meta


def cards_from_markdown_deck(md_path, source_file, systems, tag=True, llm="openrouter",
                             verbose=False, api_key=None, model=None, base_url=None):
    """Convert an authored markdown deck into store cards (source 'md').

    Existing store cards for the same source_file are matched by position and
    updated in place (same UUID, preserved tags/status/edit_history). Returns
    the list of store cards, or None if the deck has no card content."""
    system = (systems or [None])[0]
    if not system:
        return None
    title, cards = parse_markdown_deck(md_path)
    if not cards:
        return None

    meta = load_import_meta(source_file)
    store_cards = []
    for i, card in enumerate(cards):
        front, back = parse_card_content(card["content"])
        old = meta.get(i)
        c = build_card(
            front or "", back or card["content"], system, "General",
            tags=card.get("tags") or [], source="md", source_file=source_file,
            slug_hint=card["subtopic"],
            data={"original_subtopic": card["subtopic"]},
            card_id=old.get("id") if old else None,
        )
        if old:
            c["tags"] = old.get("tags") or []
            c["status"] = old.get("status") or "pending"
            c["edit_history"] = list(old.get("edit_history") or [])
        store_cards.append(c)

    if tag:
        untagged = [c for c in store_cards if not (c.get("tags") or [])]
        if untagged:
            tag_cards_with_llm(untagged, systems, llm=llm, verbose=verbose,
                               api_key=api_key, model=model, base_url=base_url)
    for c in store_cards:
        c["subtopic"] = canonical_subtopic(system, c["tags"][0]) if c["tags"] else "General"
    for c in store_cards:
        upsert_card(c)

    # Drop store cards of this source whose md section was removed
    new_ids = {c["id"] for c in store_cards}
    for old in meta.values():
        if old.get("id") not in new_ids and old.get("source") == "md":
            remove_card(old.get("id"))
    return store_cards
